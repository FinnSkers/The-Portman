# ATS-Friendly Resume Maker
from fastapi import APIRouter, HTTPException, Depends
from pydantic import BaseModel
from typing import Dict, Any, List, Optional, Union
import logging
import json
import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # Fallback if docx is not available
    Document = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
from io import BytesIO
import tempfile

router = APIRouter(prefix="/ats", tags=["ats-resume"])
logger = logging.getLogger(__name__)

# ATS Resume Models
class ATSResumeRequest(BaseModel):
    cv_data: Dict[str, Any]
    template_type: str = "clean"  # clean, minimal, professional, modern
    include_sections: List[str] = ["contact", "summary", "experience", "education", "skills"]
    keyword_optimization: bool = True
    target_job_title: Optional[str] = None
    target_industry: Optional[str] = None

class ATSResumeResponse(BaseModel):
    status: str
    resume_id: str
    format: str
    download_url: str
    preview_text: str
    ats_score: float
    optimization_suggestions: List[str]
    keyword_density: Dict[str, int]

# ATS Templates Configuration
ATS_TEMPLATES = {
    "clean": {
        "name": "Clean Professional",
        "description": "Simple, ATS-friendly format with clear sections",
        "features": ["No graphics", "Standard fonts", "Linear layout", "Keyword optimized"],
        "sections_order": ["contact", "summary", "experience", "education", "skills", "certifications"],
        "font": "Arial",
        "font_size": 11,
        "spacing": 1.15
    },
    "minimal": {
        "name": "Minimal Modern",
        "description": "Minimalist design optimized for ATS parsing",
        "features": ["Clean lines", "Maximum white space", "Easy scanning", "ATS compliant"],
        "sections_order": ["contact", "summary", "skills", "experience", "education"],
        "font": "Calibri",
        "font_size": 11,
        "spacing": 1.0
    },
    "professional": {
        "name": "Professional Standard",
        "description": "Traditional format preferred by recruiters",
        "features": ["Conservative design", "Standard layout", "Professional appearance", "ATS optimized"],
        "sections_order": ["contact", "summary", "experience", "education", "skills", "certifications", "awards"],
        "font": "Times New Roman",
        "font_size": 12,
        "spacing": 1.15
    },
    "modern": {
        "name": "Modern Clean",
        "description": "Contemporary design with ATS compatibility",
        "features": ["Modern typography", "Clean sections", "Strategic spacing", "Keyword focused"],
        "sections_order": ["contact", "summary", "skills", "experience", "education", "projects"],
        "font": "Verdana",
        "font_size": 11,
        "spacing": 1.1
    }
}

# Industry-specific keywords for optimization
INDUSTRY_KEYWORDS = {
    "software": ["agile", "scrum", "python", "javascript", "react", "node.js", "api", "database", "cloud", "devops"],
    "marketing": ["digital marketing", "seo", "social media", "content marketing", "analytics", "campaigns", "roi"],
    "finance": ["financial analysis", "budgeting", "forecasting", "excel", "financial modeling", "compliance"],
    "healthcare": ["patient care", "medical records", "hipaa", "clinical", "healthcare", "medical", "patient"],
    "education": ["curriculum", "teaching", "learning", "assessment", "classroom management", "educational"],
    "sales": ["lead generation", "client relations", "sales targets", "crm", "negotiation", "revenue"],
    "engineering": ["design", "testing", "quality assurance", "project management", "technical", "innovation"],
    "default": ["leadership", "communication", "problem solving", "teamwork", "project management", "analytical"]
}

@router.get("/templates/")
async def get_ats_templates():
    """Get available ATS resume templates."""
    return {
        "status": "success",
        "templates": ATS_TEMPLATES
    }

@router.post("/generate/", response_model=ATSResumeResponse)
async def generate_ats_resume(request: ATSResumeRequest):
    """Generate ATS-friendly resume from CV data."""
    try:
        # Validate template
        if request.template_type not in ATS_TEMPLATES:
            raise HTTPException(status_code=400, detail="Invalid template type")
        
        template = ATS_TEMPLATES[request.template_type]
        
        # Generate unique resume ID
        resume_id = f"ats_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Create DOCX resume
        doc = Document()
        
        # Configure document style
        style = doc.styles['Normal']
        font = style.font
        font.name = template["font"]
        font.size = Inches(template["font_size"] / 72)
        
        # Extract and structure CV data
        cv_data = request.cv_data
        
        # Build resume sections based on template order
        resume_content = build_ats_resume_content(cv_data, template, request)
        
        # Add sections to document
        for section_name in template["sections_order"]:
            if section_name in request.include_sections and section_name in resume_content:
                add_section_to_doc(doc, section_name, resume_content[section_name], template)
        
        # Save document
        resume_dir = "generated_resumes"
        os.makedirs(resume_dir, exist_ok=True)
        docx_path = f"{resume_dir}/{resume_id}.docx"
        doc.save(docx_path)
        
        # Generate preview text
        preview_text = generate_preview_text(resume_content)
        
        # Calculate ATS score and optimization suggestions
        ats_analysis = analyze_ats_compatibility(resume_content, request)
        
        return ATSResumeResponse(
            status="success",
            resume_id=resume_id,
            format="docx",
            download_url=f"/api/v1/ats/download/{resume_id}",
            preview_text=preview_text,
            ats_score=ats_analysis["score"],
            optimization_suggestions=ats_analysis["suggestions"],
            keyword_density=ats_analysis["keywords"]
        )
        
    except Exception as e:
        logger.error(f"ATS resume generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Resume generation failed: {str(e)}")

@router.get("/download/{resume_id}")
async def download_ats_resume(resume_id: str):
    """Download generated ATS resume."""
    try:
        docx_path = f"generated_resumes/{resume_id}.docx"
        
        if not os.path.exists(docx_path):
            raise HTTPException(status_code=404, detail="Resume not found")
        
        from fastapi.responses import FileResponse
        return FileResponse(
            path=docx_path,
            filename=f"{resume_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"Resume download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@router.post("/analyze/")
async def analyze_ats_score(cv_data: Dict[str, Any], target_job_title: Optional[str] = None):
    """Analyze ATS compatibility score for CV data."""
    try:
        # Determine industry keywords
        industry = determine_industry(cv_data, target_job_title)
        keywords = INDUSTRY_KEYWORDS.get(industry, INDUSTRY_KEYWORDS["default"])
        
        # Analyze content
        analysis = analyze_ats_compatibility(cv_data, None, keywords)
        
        return {
            "status": "success",
            "ats_score": analysis["score"],
            "industry": industry,
            "keyword_matches": analysis["keywords"],
            "suggestions": analysis["suggestions"],
            "missing_keywords": analysis.get("missing_keywords", [])
        }
        
    except Exception as e:
        logger.error(f"ATS analysis failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

def build_ats_resume_content(cv_data: Dict[str, Any], template: Dict[str, Any], request: ATSResumeRequest) -> Dict[str, Any]:
    """Build structured resume content from CV data."""
    content = {}
    
    # Contact Information
    if "contact" in request.include_sections:
        contact = {
            "name": cv_data.get("name", ""),
            "email": cv_data.get("email", ""),
            "phone": cv_data.get("phone", cv_data.get("phoneNumber", "")),
            "location": format_address(cv_data.get("address", {})),
            "linkedin": extract_linkedin(cv_data.get("links", [])),
            "website": extract_website(cv_data.get("links", []))
        }
        content["contact"] = contact
    
    # Professional Summary
    if "summary" in request.include_sections:
        summary = cv_data.get("summary", "")
        if request.keyword_optimization and request.target_industry:
            summary = optimize_summary_keywords(summary, request.target_industry)
        content["summary"] = summary
    
    # Experience
    if "experience" in request.include_sections:
        experience = format_experience_for_ats(cv_data.get("experience", []), request)
        content["experience"] = experience
    
    # Education
    if "education" in request.include_sections:
        education = format_education_for_ats(cv_data.get("education", []))
        content["education"] = education
    
    # Skills
    if "skills" in request.include_sections:
        skills = format_skills_for_ats(cv_data.get("skills", []), request)
        content["skills"] = skills
    
    # Certifications
    if "certifications" in request.include_sections:
        certifications = format_certifications_for_ats(cv_data.get("certifications", []))
        content["certifications"] = certifications
    
    # Projects
    if "projects" in request.include_sections:
        projects = format_projects_for_ats(cv_data.get("projects", []))
        content["projects"] = projects
    
    return content

def add_section_to_doc(doc: Document, section_name: str, section_data: Any, template: Dict[str, Any]):
    """Add a section to the Word document."""
    # Section headers mapping
    headers = {
        "contact": "",  # No header for contact
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "skills": "CORE COMPETENCIES",
        "certifications": "CERTIFICATIONS",
        "projects": "KEY PROJECTS",
        "awards": "AWARDS & RECOGNITION"
    }
    
    if section_name == "contact":
        # Add contact information at top
        if isinstance(section_data, dict):
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(section_data.get("name", ""))
            name_run.bold = True
            name_run.font.size = Inches(16 / 72)
            
            contact_info = []
            if section_data.get("email"):
                contact_info.append(section_data["email"])
            if section_data.get("phone"):
                contact_info.append(section_data["phone"])
            if section_data.get("location"):
                contact_info.append(section_data["location"])
            
            if contact_info:
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.add_run(" | ".join(contact_info))
            
            if section_data.get("linkedin") or section_data.get("website"):
                links_info = []
                if section_data.get("linkedin"):
                    links_info.append(f"LinkedIn: {section_data['linkedin']}")
                if section_data.get("website"):
                    links_info.append(f"Website: {section_data['website']}")
                
                links_para = doc.add_paragraph()
                links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                links_para.add_run(" | ".join(links_info))
    
    else:
        # Add section header
        if headers.get(section_name):
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(headers[section_name])
            header_run.bold = True
            header_run.font.size = Inches(12 / 72)
        
        # Add section content based on type
        if section_name == "summary" and isinstance(section_data, str):
            doc.add_paragraph(section_data)
        
        elif section_name == "experience" and isinstance(section_data, list):
            for exp in section_data:
                # Job title and company
                job_para = doc.add_paragraph()
                title_run = job_para.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                title_run.bold = True
                
                # Duration and location
                if exp.get('duration') or exp.get('location'):
                    duration_para = doc.add_paragraph()
                    duration_info = []
                    if exp.get('duration'):
                        duration_info.append(exp['duration'])
                    if exp.get('location'):
                        duration_info.append(exp['location'])
                    duration_para.add_run(" | ".join(duration_info))
                
                # Responsibilities
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities'][:5]:  # Limit to 5 bullets
                        bullet_para = doc.add_paragraph(f"• {resp}", style='List Bullet')
        
        elif section_name == "education" and isinstance(section_data, list):
            for edu in section_data:
                edu_para = doc.add_paragraph()
                edu_text = f"{edu.get('degree', '')} in {edu.get('field', '')}"
                if edu.get('institution'):
                    edu_text += f" - {edu['institution']}"
                if edu.get('year'):
                    edu_text += f" ({edu['year']})"
                edu_para.add_run(edu_text)
        
        elif section_name == "skills" and isinstance(section_data, dict):
            if section_data.get('technical'):
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technical Skills: ").bold = True
                tech_para.add_run(", ".join(section_data['technical']))
            
            if section_data.get('soft'):
                soft_para = doc.add_paragraph()
                soft_para.add_run("Soft Skills: ").bold = True
                soft_para.add_run(", ".join(section_data['soft']))
    
    # Add spacing between sections
    doc.add_paragraph()

def format_experience_for_ats(experience: List[Dict], request: ATSResumeRequest) -> List[Dict]:
    """Format experience data for ATS compatibility."""
    formatted_exp = []
    
    for exp in experience[:6]:  # Limit to 6 most recent positions
        formatted = {
            "title": exp.get("title", exp.get("position", "")),
            "company": exp.get("company", exp.get("organization", "")),
            "duration": exp.get("duration", ""),
            "location": exp.get("location", ""),
            "responsibilities": []
        }
        
        # Extract responsibilities
        if exp.get("responsibilities"):
            if isinstance(exp["responsibilities"], list):
                for resp in exp["responsibilities"][:5]:
                    if isinstance(resp, dict):
                        task = resp.get("task", "")
                        desc = resp.get("description", "")
                        if task and desc:
                            formatted["responsibilities"].append(f"{task}: {desc}")
                        elif task:
                            formatted["responsibilities"].append(task)
                    else:
                        formatted["responsibilities"].append(str(resp))
            else:
                formatted["responsibilities"].append(str(exp["responsibilities"]))
        
        # Add keyword optimization
        if request.keyword_optimization and request.target_industry:
            keywords = INDUSTRY_KEYWORDS.get(request.target_industry, INDUSTRY_KEYWORDS["default"])
            formatted["responsibilities"] = optimize_responsibilities_keywords(
                formatted["responsibilities"], keywords
            )
        
        formatted_exp.append(formatted)
    
    return formatted_exp

def format_education_for_ats(education: List[Dict]) -> List[Dict]:
    """Format education data for ATS compatibility."""
    formatted_edu = []
    
    for edu in education:
        formatted = {
            "degree": edu.get("degree", ""),
            "field": edu.get("field", ""),
            "institution": edu.get("institution", ""),
            "year": edu.get("year", ""),
            "gpa": edu.get("gpa", ""),
            "achievements": edu.get("achievements", [])
        }
        formatted_edu.append(formatted)
    
    return formatted_edu

def format_skills_for_ats(skills: List, request: ATSResumeRequest) -> Dict[str, List]:
    """Format skills data for ATS compatibility."""
    # Convert skills to standardized format
    skill_list = []
    if isinstance(skills, list):
        for skill in skills:
            if isinstance(skill, dict):
                skill_list.append(skill.get("skill", skill.get("name", "")))
            else:
                skill_list.append(str(skill))
    
    # Categorize skills
    technical_keywords = ["python", "javascript", "react", "node", "sql", "aws", "docker", "git"]
    soft_keywords = ["leadership", "communication", "management", "teamwork", "problem"]
    
    technical_skills = []
    soft_skills = []
    
    for skill in skill_list:
        skill_lower = skill.lower()
        if any(tech in skill_lower for tech in technical_keywords):
            technical_skills.append(skill)
        elif any(soft in skill_lower for soft in soft_keywords):
            soft_skills.append(skill)
        else:
            technical_skills.append(skill)  # Default to technical
    
    # Add industry-specific keywords if optimization enabled
    if request.keyword_optimization and request.target_industry:
        industry_keywords = INDUSTRY_KEYWORDS.get(request.target_industry, [])
        for keyword in industry_keywords:
            if keyword not in [s.lower() for s in technical_skills]:
                technical_skills.append(keyword.title())
    
    return {
        "technical": technical_skills[:15],  # Limit to 15 technical skills
        "soft": soft_skills[:8]  # Limit to 8 soft skills
    }

def format_certifications_for_ats(certifications: List) -> List[Dict]:
    """Format certifications for ATS compatibility."""
    formatted_certs = []
    
    for cert in certifications:
        if isinstance(cert, dict):
            formatted_certs.append({
                "name": cert.get("name", ""),
                "issuer": cert.get("issuer", ""),
                "date": cert.get("date", ""),
                "expiry": cert.get("expiry", "")
            })
        else:
            formatted_certs.append({"name": str(cert)})
    
    return formatted_certs

def format_projects_for_ats(projects: List) -> List[Dict]:
    """Format projects for ATS compatibility."""
    formatted_projects = []
    
    for proj in projects[:4]:  # Limit to 4 projects
        if isinstance(proj, dict):
            formatted_projects.append({
                "name": proj.get("name", ""),
                "description": proj.get("description", ""),
                "technologies": proj.get("technologies", []),
                "achievements": proj.get("achievements", [])
            })
    
    return formatted_projects

def analyze_ats_compatibility(content: Dict[str, Any], request: Optional[ATSResumeRequest], keywords: Optional[List[str]] = None) -> Dict[str, Any]:
    """Analyze ATS compatibility and provide optimization suggestions."""
    score = 0.0
    suggestions = []
    keyword_matches = {}
    
    # Determine keywords to check
    if keywords is None and request and request.target_industry:
        keywords = INDUSTRY_KEYWORDS.get(request.target_industry, INDUSTRY_KEYWORDS["default"])
    elif keywords is None:
        keywords = INDUSTRY_KEYWORDS["default"]
    
    # Check content structure (30% of score)
    structure_score = 0
    if content.get("contact"):
        structure_score += 10
    if content.get("summary"):
        structure_score += 5
    if content.get("experience"):
        structure_score += 10
    if content.get("education"):
        structure_score += 3
    if content.get("skills"):
        structure_score += 2
    
    score += min(structure_score, 30)
    
    # Check keyword density (40% of score)
    text_content = json.dumps(content).lower()
    keyword_count = 0
    for keyword in keywords:
        count = text_content.count(keyword.lower())
        if count > 0:
            keyword_matches[keyword] = count
            keyword_count += min(count, 3)  # Cap at 3 mentions per keyword
    
    keyword_score = min((keyword_count / len(keywords)) * 40, 40)
    score += keyword_score
    
    # Check formatting (30% of score)
    format_score = 30  # Assume good formatting for generated resumes
    score += format_score
    
    # Generate suggestions
    if score < 60:
        suggestions.append("Consider adding more relevant keywords to improve ATS compatibility")
    if not content.get("summary"):
        suggestions.append("Add a professional summary to highlight your key qualifications")
    if len(keyword_matches) < len(keywords) * 0.3:
        suggestions.append("Include more industry-specific keywords and skills")
    if content.get("experience") and len(content["experience"]) < 2:
        suggestions.append("Expand your experience section with more detailed descriptions")
    
    return {
        "score": round(score, 1),
        "suggestions": suggestions,
        "keywords": keyword_matches,
        "missing_keywords": [k for k in keywords if k not in keyword_matches]
    }

def generate_preview_text(content: Dict[str, Any]) -> str:
    """Generate preview text from resume content."""
    preview_parts = []
    
    if content.get("contact", {}).get("name"):
        preview_parts.append(f"Name: {content['contact']['name']}")
    
    if content.get("summary"):
        preview_parts.append(f"Summary: {content['summary'][:100]}...")
    
    if content.get("experience"):
        exp_count = len(content["experience"])
        preview_parts.append(f"Experience: {exp_count} position(s)")
    
    if content.get("skills"):
        skill_count = len(content["skills"].get("technical", [])) + len(content["skills"].get("soft", []))
        preview_parts.append(f"Skills: {skill_count} listed")
    
    return " | ".join(preview_parts)

# Helper functions
def format_address(address: Dict) -> str:
    """Format address for ATS compatibility."""
    if isinstance(address, str):
        return address
    
    parts = []
    if address.get("city"):
        parts.append(address["city"])
    if address.get("state"):
        parts.append(address["state"])
    if address.get("country"):
        parts.append(address["country"])
    
    return ", ".join(parts)

def extract_linkedin(links: List) -> str:
    """Extract LinkedIn URL from links."""
    for link in links:
        if isinstance(link, dict):
            url = link.get("url", "")
            if "linkedin" in url.lower():
                return url
        elif isinstance(link, str) and "linkedin" in link.lower():
            return link
    return ""

def extract_website(links: List) -> str:
    """Extract website URL from links."""
    for link in links:
        if isinstance(link, dict):
            url = link.get("url", "")
            if "linkedin" not in url.lower() and url.startswith(("http", "www")):
                return url
        elif isinstance(link, str) and "linkedin" not in link.lower() and link.startswith(("http", "www")):
            return link
    return ""

def determine_industry(cv_data: Dict[str, Any], target_job_title: Optional[str] = None) -> str:
    """Determine industry from CV data."""
    if target_job_title:
        title_lower = target_job_title.lower()
        if any(word in title_lower for word in ["software", "developer", "engineer", "programmer"]):
            return "software"
        elif any(word in title_lower for word in ["marketing", "digital", "social media"]):
            return "marketing"
        elif any(word in title_lower for word in ["finance", "financial", "accounting"]):
            return "finance"
        elif any(word in title_lower for word in ["sales", "account", "business development"]):
            return "sales"
    
    # Analyze skills and experience
    text_content = json.dumps(cv_data).lower()
    if any(word in text_content for word in ["python", "javascript", "software", "programming"]):
        return "software"
    elif any(word in text_content for word in ["marketing", "seo", "digital"]):
        return "marketing"
    elif any(word in text_content for word in ["finance", "accounting", "financial"]):
        return "finance"
    
    return "default"

def optimize_summary_keywords(summary: str, industry: str) -> str:
    """Optimize summary with industry keywords."""
    keywords = INDUSTRY_KEYWORDS.get(industry, INDUSTRY_KEYWORDS["default"])
    # Simple keyword injection (could be enhanced with AI)
    for keyword in keywords[:3]:
        if keyword.lower() not in summary.lower():
            summary += f" Experienced in {keyword}."
    return summary

def optimize_responsibilities_keywords(responsibilities: List[str], keywords: List[str]) -> List[str]:
    """Optimize responsibilities with keywords."""
    optimized = []
    for resp in responsibilities:
        resp_lower = resp.lower()
        for keyword in keywords:
            if keyword.lower() not in resp_lower and len(optimized) < len(responsibilities):
                resp = resp.replace(".", f" using {keyword}.")
                break
        optimized.append(resp)
    return optimized
